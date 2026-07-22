#!/usr/bin/env python3
"""Build the interview speaker-notes DOCX from docs/SPEAKER_NOTES.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SPEAKER_NOTES.md"
OUTPUT = ROOT / "public" / "downloads" / "cloudflare-application-security-speaker-notes.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GREEN = "4E7C20"
MUTED = "66736A"
LIGHT_GREEN = "EFF8E7"
LIGHT_BLUE = "E8EEF5"
LIGHT_ORANGE = "FFF4E5"
ORANGE = "B96612"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_paragraph_shading(paragraph, fill: str, border_color: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://[^\s]+)")


def add_inline(paragraph, text: str):
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True, color=INK)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            trailing = ""
            while token and token[-1] in ".,);":
                trailing = token[-1] + trailing
                token = token[:-1]
            add_hyperlink(paragraph, token, token)
            if trailing:
                set_run_font(paragraph.add_run(trailing))
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]))


def next_id(elements, tag: str, attribute: str) -> int:
    values = [int(element.get(qn(attribute))) for element in elements.findall(qn(tag))]
    return max(values, default=0) + 1


def create_abstract_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_id(numbering, "w:abstractNum", "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, indent])
    level.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_id = next_id(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_list_paragraph(doc, text: str, num_id: int, recovery=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    add_inline(paragraph, text)
    if recovery:
        set_paragraph_shading(paragraph, LIGHT_ORANGE, ORANGE)
    return paragraph


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.add_run("CLOUDFLARE APPLICATION SECURITY DEMO  ·  SPEAKER NOTES"), size=8, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    set_run_font(footer_p.add_run("INTERVIEW EDITION  ·  "), size=8, color=MUTED)
    add_page_field(footer_p)


def add_cover(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(kicker.add_run("PRESENTER RUNBOOK"), size=10, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    set_run_font(title.add_run("Application Security\nat the Edge"), size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("Speaker Notes — Candidate-Neutral Interview Edition"), size=13.5, color=MUTED)

    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=120)
    values = [
        ("12–15 MIN", "Core story"),
        ("14 SLIDES", "Guided sequence"),
        ("7 CHECKS", "Live preflight"),
        ("200 → 403", "Proof moment"),
    ]
    for cell, (metric, label) in zip(table.rows[0].cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(metric), size=12, color=GREEN, bold=True)
        p.add_run("\n")
        set_run_font(p.add_run(label), size=8.5, color=MUTED)

    thesis = doc.add_paragraph()
    thesis.paragraph_format.space_before = Pt(28)
    thesis.paragraph_format.space_after = Pt(8)
    thesis.paragraph_format.left_indent = Inches(0.18)
    thesis.paragraph_format.right_indent = Inches(0.18)
    set_paragraph_shading(thesis, LIGHT_BLUE, BLUE)
    set_run_font(thesis.add_run("CORE THESIS  "), size=9, color=BLUE, bold=True)
    set_run_font(
        thesis.add_run("Cloudflare makes the network edge the first enforcement point, before application code executes."),
        size=12,
        color=INK,
        bold=True,
    )

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(28)
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(meta.add_run("Production deck  "), size=9, color=MUTED, bold=True)
    add_hyperlink(meta, "innovativefuturesolutions.com", "https://innovativefuturesolutions.com")


def parse_markdown(
    doc: Document,
    markdown: str,
    bullet_abstract: int,
    number_abstract: int,
    *,
    page_break_first_heading: bool = False,
):
    current_heading = ""
    current_list_kind = None
    current_num_id = None
    first_heading_seen = False

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            current_list_kind = None
            current_num_id = None
            continue
        if line == "---":
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Heading 1")
            if page_break_first_heading and not first_heading_seen:
                paragraph.paragraph_format.page_break_before = True
            add_inline(paragraph, line[2:])
            current_heading = line[2:]
            current_list_kind = None
            first_heading_seen = True
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, line[3:])
            current_heading = line[3:]
            current_list_kind = None
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline(paragraph, line[4:])
            current_heading = line[4:]
            current_list_kind = None
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        bullet = re.match(r"^-\s+(.*)$", line)
        if numbered or bullet:
            kind = "number" if numbered else "bullet"
            if current_list_kind != kind or current_num_id is None:
                current_num_id = create_num_instance(doc, number_abstract if kind == "number" else bullet_abstract)
                current_list_kind = kind
            add_list_paragraph(
                doc,
                (numbered or bullet).group(1),
                current_num_id,
                recovery=current_heading == "Recovery",
            )
            continue

        paragraph = doc.add_paragraph()
        if line.startswith("**Timing:**") or line.startswith("**Objective:**") or line.startswith("**Production deck:**") or line.startswith("**Target duration:**") or line.startswith("**Core thesis:**"):
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
        if current_heading == "Recovery":
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.right_indent = Inches(0.1)
            set_paragraph_shading(paragraph, LIGHT_ORANGE, ORANGE)
        add_inline(paragraph, line.replace("  ", ""))


def build():
    source = SOURCE.read_text(encoding="utf-8")
    blocks = [block.strip() for block in source.split("\n---\n") if block.strip()]

    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    bullet_abstract = create_abstract_numbering(doc, bullet=True)
    number_abstract = create_abstract_numbering(doc, bullet=False)

    doc.core_properties.title = "Application Security at the Edge — Speaker Notes"
    doc.core_properties.subject = "Candidate-neutral Cloudflare application security interview runbook"
    doc.core_properties.author = "Innovative Future Solutions"
    doc.core_properties.keywords = "Cloudflare, application security, speaker notes, interview"

    add_cover(doc)
    doc.add_page_break()

    intro_lines = blocks[0].splitlines()
    start = next(index for index, line in enumerate(intro_lines) if line == "## Before the interview")
    intro = "# Presenter preflight and run of show\n\n" + "\n".join(intro_lines[start:])
    parse_markdown(doc, intro, bullet_abstract, number_abstract)

    for block in blocks[1:]:
        parse_markdown(
            doc,
            block,
            bullet_abstract,
            number_abstract,
            page_break_first_heading=True,
        )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
